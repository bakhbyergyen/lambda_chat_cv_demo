from datetime import datetime
from docx import Document # type: ignore
from docx.shared import  Pt, Cm, RGBColor # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT # type: ignore
from docx.enum.style import WD_STYLE_TYPE # type: ignore
from docx.oxml.ns import qn # type: ignore
from datetime import datetime 
from docx.oxml import OxmlElement # type: ignore

def add_horizontal_line(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_horizontal_line(p)

def generate_docx_cv(cv_data):
    doc = Document()

    # set page size to A4 and margins to narrow
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)

    
    # Set up styles
    styles = doc.styles

    # Name style
    name_style = styles.add_style('Name', WD_STYLE_TYPE.PARAGRAPH)
    name_style.font.name = 'Roboto'
    name_style.font.size = Pt(16)
    name_style.font.bold = True
    name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_style.paragraph_format.space_after = Pt(10)

    # Title style
    position_style = styles.add_style('Position', WD_STYLE_TYPE.PARAGRAPH)
    position_style.font.name = 'Roboto'
    position_style.font.size = Pt(16)
    position_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    position_style.paragraph_format.space_after = Pt(10)
    # Contact style
    contact_style = styles.add_style('Contact', WD_STYLE_TYPE.PARAGRAPH)
    contact_style.font.name = 'Roboto'
    contact_style.font.size = Pt(9)
    contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_style.paragraph_format.space_after = Pt(12)

    # Section Header style
    section_header_style = styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
    section_header_style.font.name = 'Roboto'
    section_header_style.font.size = Pt(12)
    section_header_style.font.bold = True
    section_header_style.font.color.rgb = RGBColor(139, 0, 0)  # Dark red
    section_header_style.paragraph_format.space_before = Pt(12)
    section_header_style.paragraph_format.space_after = Pt(6)

    # Company Name style
    company_name_style = styles.add_style('CompanyName', WD_STYLE_TYPE.PARAGRAPH)
    company_name_style.font.name = 'Roboto'
    company_name_style.font.size = Pt(11)
    company_name_style.font.bold = True

    # Job Title style
    job_title_style = styles.add_style('JobTitle', WD_STYLE_TYPE.PARAGRAPH)
    job_title_style.font.name = 'Roboto'
    job_title_style.font.size = Pt(11)

    # Date Location style
    date_location_style = styles.add_style('DateLocation', WD_STYLE_TYPE.PARAGRAPH)
    date_location_style.font.name = 'Roboto'
    date_location_style.font.size = Pt(9)
    date_location_style.font.color.rgb = RGBColor(128, 128, 128)  # Gray

    # Normal style
    normal_style = styles['Normal']
    normal_style.font.name = 'Roboto'
    normal_style.font.size = Pt(10)

    # Add personal information
    doc.add_paragraph(cv_data["personal_info"]["name"], style='Name')
    doc.add_paragraph(cv_data["personal_info"]["position"], style='Position')
    contact_info = f"{cv_data['personal_info']['email']} | {cv_data['personal_info']['phone']} | {cv_data['personal_info']['location']}"
    if 'linkedin' in cv_data['personal_info']:
        contact_info += f" | {cv_data['personal_info']['linkedin']}"
    doc.add_paragraph(contact_info, style='Contact')

    # Add divider
    add_divider(doc)

    # Work Experience
    doc.add_paragraph("WORK EXPERIENCE", style='SectionHeader')
    for job in cv_data["work_experience"]:
        p = doc.add_paragraph(style='CompanyName')
        p.add_run(f"{job['company']} - ").bold = True
        p.add_run(job['position']).bold = False
        p.paragraph_format.space_before = Pt(0)
        p = doc.add_paragraph(style='DateLocation')
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(19), WD_TAB_ALIGNMENT.RIGHT)
        p.paragraph_format.space_before = Pt(0)  # Remove space before this paragraph
        p.add_run(f"{job['start_date']} - {job['end_date']}")
        p.add_run('\t' + job['location'])

        for responsibility in job["responsibilities"]:
            doc.add_paragraph(f"• {responsibility}", style='Normal')

    # Add divider
    add_divider(doc)
    # Education
    doc.add_paragraph("EDUCATION", style='SectionHeader')
    for edu in cv_data["education"]:
        p = doc.add_paragraph(style='CompanyName')
        p.add_run(f"{edu['institution']} - ").bold = True
        p.add_run(f"{edu['degree']}, {edu['field']}").bold = False
        
        p = doc.add_paragraph(style='DateLocation')
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(19), WD_TAB_ALIGNMENT.RIGHT)
        p.paragraph_format.space_before = Pt(0)  # Remove space before this paragraph
        p.add_run(f"{edu['start_date']} - {edu['end_date']}")
        p.add_run(f"\t{edu['location']}")

    # Add divider
    add_divider(doc)

    # Skills
    doc.add_paragraph("SKILLS", style='SectionHeader')
    skills_text = "; ".join(cv_data["skills"])
    doc.add_paragraph(skills_text, style='Normal')

    # Hobbies (if provided)
    if "hobbies" in cv_data and cv_data["hobbies"]:
        doc.add_paragraph("HOBBIES", style='SectionHeader')
        hobbies_text = "; ".join(cv_data["hobbies"])
        doc.add_paragraph(hobbies_text, style='Normal')

    # Save the document
    filename = f"{cv_data['personal_info']['name'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    return filename